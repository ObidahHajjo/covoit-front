from __future__ import annotations

import json
import re
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Callable
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageOps


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
SCREENSHOTS_DIR = DOCS_DIR / "screenshots"
I18N_FILE = ROOT / "src" / "i18n" / "dictionaries.ts"
CSS_FILE = ROOT / "src" / "index.css"
PACKAGE_FILE = ROOT / "package.json"


@dataclass
class BrandTokens:
    primary: str
    primary_soft: str
    ink: str
    muted: str
    body_font: str
    heading_font: str


@dataclass
class StepGroup:
    heading: str
    instructions: list[str]
    screenshot_path: Path | None
    placeholder: str | None
    caption: str | None


@dataclass
class Chapter:
    title: str
    overview: str
    steps: list[StepGroup]
    notes: str


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def parse_dictionaries() -> dict[str, dict[str, str]]:
    dictionaries: dict[str, dict[str, str]] = {}
    current_locale: str | None = None
    for line in I18N_FILE.read_text(encoding="utf-8").splitlines():
        locale_match = re.match(r"\s{2}(en|fr|ar): \{$", line)
        if locale_match:
            current_locale = locale_match.group(1)
            dictionaries[current_locale] = {}
            continue

        if current_locale and re.match(r"\s{2}\},?$", line):
            current_locale = None
            continue

        if not current_locale:
            continue

        entry_match = re.match(r'\s{4}"([^"]+)":\s*"((?:\\.|[^"\\])*)",?$', line)
        if not entry_match:
            continue

        key = entry_match.group(1)
        value = json.loads(f'"{entry_match.group(2)}"')
        dictionaries[current_locale][key] = value

    return dictionaries


def parse_brand_tokens() -> BrandTokens:
    css = CSS_FILE.read_text(encoding="utf-8")

    def extract(name: str) -> str:
        match = re.search(rf"{re.escape(name)}:\s*([^;]+);", css)
        if not match:
            raise ValueError(f"Missing token: {name}")
        return match.group(1).strip().strip('"')

    return BrandTokens(
        primary=extract("--theme-primary"),
        primary_soft=extract("--theme-primary-soft"),
        ink=extract("--theme-ink"),
        muted=extract("--theme-muted"),
        body_font=extract("--font-body").split(",")[0].strip('"'),
        heading_font=extract("--font-heading").split(",")[0].strip('"'),
    )


def hex_to_rgb(value: str) -> RGBColor:
    value = value.lstrip("#")
    return RGBColor.from_string(value.upper())


def tr_factory(translations: dict[str, str]) -> Callable[[str, dict[str, str | int] | None], str]:
    def tr(key: str, values: dict[str, str | int] | None = None) -> str:
        message = translations[key]
        for name, value in (values or {}).items():
            message = message.replace(f"{{{name}}}", str(value))
        return message

    return tr


def build_cover(doc: Document, locale: str, app_name: str, version: str) -> None:
    title = "Guide utilisateur" if locale == "fr" else "User guide"
    generated_label = "Date" if locale == "fr" else "Date"
    logo_label = "[LOGO PLACEHOLDER: logo Covoit]" if locale == "fr" else "[LOGO PLACEHOLDER: Covoit logo]"

    accent = doc.add_paragraph(style="Intense Quote")
    accent.alignment = WD_ALIGN_PARAGRAPH.CENTER
    accent.add_run(" ")

    logo = doc.add_paragraph(style="Caption")
    logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo.add_run(logo_label)

    title_paragraph = doc.add_paragraph(style="Title")
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.add_run(app_name)

    subtitle = doc.add_paragraph(style="Heading 2")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(title)

    version_paragraph = doc.add_paragraph(style="Normal")
    version_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version_paragraph.add_run(f"Version {version}")

    date_paragraph = doc.add_paragraph(style="Normal")
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_paragraph.add_run(f"{generated_label}: {date.today().isoformat()}")

    doc.add_page_break()


def add_toc(doc: Document, locale: str) -> None:
    heading = doc.add_paragraph(style="Heading 1")
    heading.add_run("Table des matieres" if locale == "fr" else "Table of contents")

    paragraph = doc.add_paragraph(style="Normal")
    run = paragraph.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:r")
    placeholder_text = OxmlElement("w:t")
    placeholder_text.text = "Update the table of contents in Word if page numbers are not displayed." if locale == "en" else "Mettez a jour la table des matieres dans Word si les numeros de page ne s'affichent pas."
    placeholder.append(placeholder_text)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_separate)
    run._r.append(placeholder)
    run._r.append(fld_end)

    doc.add_page_break()


def add_heading_border(paragraph) -> None:
    paragraph_format = paragraph._element.get_or_add_pPr()
    borders = paragraph_format.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        paragraph_format.append(borders)

    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)

    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), "526448")


def ensure_bordered_image(path: Path, border_color: str) -> None:
    image = Image.open(path)
    bordered = ImageOps.expand(image, border=3, fill=border_color)
    bordered.save(path)


def add_screenshot(doc: Document, step: StepGroup, image_width_cm: float = 15.5) -> tuple[int, int]:
    if step.screenshot_path and step.screenshot_path.exists():
        picture_paragraph = doc.add_paragraph(style="Normal")
        picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture_paragraph.add_run().add_picture(str(step.screenshot_path), width=Cm(image_width_cm))
        caption = doc.add_paragraph(style="Caption")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.add_run(step.caption or step.screenshot_path.name)
        return (1, 0)

    placeholder = doc.add_paragraph(style="Caption")
    placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
    placeholder.add_run(step.placeholder or "[SCREENSHOT: Missing screenshot]")
    return (0, 1)


def add_numbered_list(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.add_run(item)


def add_chapter(doc: Document, chapter: Chapter) -> tuple[int, int]:
    heading = doc.add_paragraph(style="Heading 1")
    heading.add_run(chapter.title)
    add_heading_border(heading)

    overview = doc.add_paragraph(style="Normal")
    overview.add_run(chapter.overview)

    embedded = 0
    placeholders = 0
    for index, step in enumerate(chapter.steps, start=1):
        subheading = doc.add_paragraph(style="Heading 2")
        subheading.add_run(f"{index}. {step.heading}")
        add_numbered_list(doc, step.instructions)
        embedded_count, placeholder_count = add_screenshot(doc, step)
        embedded += embedded_count
        placeholders += placeholder_count

    notes = doc.add_paragraph(style="Intense Quote")
    notes.add_run(chapter.notes)
    return embedded, placeholders


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill.replace("#", ""))


def customize_styles(doc: Document, brand: BrandTokens) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = brand.body_font
    normal.font.size = Pt(11)
    normal.font.color.rgb = hex_to_rgb(brand.ink)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    title = styles["Title"]
    title.font.name = brand.heading_font
    title.font.size = Pt(28)
    title.font.bold = True
    title.font.color.rgb = hex_to_rgb(brand.primary)
    title.paragraph_format.space_after = Pt(12)

    heading1 = styles["Heading 1"]
    heading1.font.name = brand.heading_font
    heading1.font.size = Pt(20)
    heading1.font.bold = True
    heading1.font.color.rgb = hex_to_rgb(brand.ink)
    heading1.paragraph_format.space_before = Pt(8)
    heading1.paragraph_format.space_after = Pt(8)

    heading2 = styles["Heading 2"]
    heading2.font.name = brand.heading_font
    heading2.font.size = Pt(14)
    heading2.font.bold = True
    heading2.font.color.rgb = hex_to_rgb(brand.primary)
    heading2.paragraph_format.space_before = Pt(8)
    heading2.paragraph_format.space_after = Pt(4)

    heading3 = styles["Heading 3"]
    heading3.font.name = brand.heading_font
    heading3.font.size = Pt(12)
    heading3.font.bold = True
    heading3.font.color.rgb = hex_to_rgb(brand.muted)

    caption = styles["Caption"]
    caption.font.name = brand.body_font
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor(112, 119, 116)
    caption.paragraph_format.space_after = Pt(8)

    intense_quote = styles["Intense Quote"]
    intense_quote.font.name = brand.body_font
    intense_quote.font.size = Pt(10.5)
    intense_quote.font.color.rgb = hex_to_rgb(brand.ink)
    intense_quote.paragraph_format.space_before = Pt(8)
    intense_quote.paragraph_format.space_after = Pt(8)
    intense_quote.paragraph_format.left_indent = Cm(0)
    intense_quote.paragraph_format.right_indent = Cm(0)

    intense_ppr = intense_quote.element.get_or_add_pPr()
    shading = intense_ppr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        intense_ppr.append(shading)
    shading.set(qn("w:fill"), brand.primary_soft.replace("#", ""))

    list_number = styles["List Number"]
    list_number.font.name = brand.body_font
    list_number.font.size = Pt(11)
    list_number.font.color.rgb = hex_to_rgb(brand.ink)
    list_number.paragraph_format.space_after = Pt(4)
    list_number.paragraph_format.line_spacing = 1.15


def fr_chapters(tr: Callable[[str, dict[str, str | int] | None], str]) -> list[Chapter]:
    return [
        Chapter(
            title="Inscription et acces",
            overview="Utilisez les ecrans de connexion pour ouvrir votre espace Covoit, creer un compte ou recuperer l'acces a votre profil.",
            steps=[
                StepGroup(
                    heading="Se connecter",
                    instructions=[
                        f"Ouvrez la page {tr('auth.signIn')}.",
                        f"Saisissez {tr('auth.emailAddress')} puis {tr('auth.password')}.",
                        f"Cliquez sur {tr('auth.signIn')} pour acceder a votre espace.",
                    ],
                    screenshot_path=SCREENSHOTS_DIR / "screenshot-authentication-1.png",
                    placeholder=None,
                    caption="Figure 1 - Ecran de connexion",
                ),
                StepGroup(
                    heading="Creer un compte",
                    instructions=[
                        f"Depuis la connexion, cliquez sur {tr('auth.createAccount')}.",
                        f"Renseignez {tr('common.email')}, {tr('common.password')} et {tr('common.confirm')}.",
                        f"Validez avec {tr('auth.createAccount')} pour passer a l'etape de profil.",
                    ],
                    screenshot_path=SCREENSHOTS_DIR / "screenshot-authentication-2.png",
                    placeholder=None,
                    caption="Figure 2 - Ecran de creation de compte",
                ),
                StepGroup(
                    heading="Demander un lien de reinitialisation",
                    instructions=[
                        f"Ouvrez la page {tr('auth.forgotPassword')}.",
                        f"Saisissez votre {tr('common.email')}.",
                        f"Cliquez sur {tr('auth.sendResetLinkButton')} pour recevoir les instructions.",
                    ],
                    screenshot_path=SCREENSHOTS_DIR / "screenshot-authentication-3.png",
                    placeholder=None,
                    caption="Figure 3 - Ecran d'envoi du lien de reinitialisation",
                ),
                StepGroup(
                    heading="Definir un nouveau mot de passe",
                    instructions=[
                        f"Ouvrez le lien recu puis verifiez le champ {tr('common.email')}.",
                        f"Saisissez {tr('auth.newPassword')} et {tr('auth.confirmPassword')}.",
                        f"Cliquez sur {tr('auth.resetPassword')} pour terminer la mise a jour.",
                    ],
                    screenshot_path=SCREENSHOTS_DIR / "screenshot-authentication-4.png",
                    placeholder=None,
                    caption="Figure 4 - Ecran de reinitialisation du mot de passe",
                ),
            ],
            notes="Astuce - Gardez le meme e-mail sur tous vos trajets pour retrouver plus facilement vos reservations et vos conversations.",
        ),
        Chapter(
            title="Completer le profil",
            overview="Apres l'inscription, completez votre identite de base pour debloquer les pages protegees et apparaitre clairement dans la communaute.",
            steps=[
                StepGroup(
                    heading="Verifier l'e-mail de reference",
                    instructions=[
                        f"Ouvrez l'ecran {tr('profile.completeProfile')}.",
                        f"Controlez le champ {tr('common.email')} affiche en reference.",
                        "Confirmez que vous terminez le profil du bon compte avant de continuer.",
                    ],
                    screenshot_path=None,
                    placeholder="[SCREENSHOT: Ecran Completer le profil avec le champ E-mail visible - capture manuelle requise apres authentification]",
                    caption=None,
                ),
                StepGroup(
                    heading="Renseigner l'identite",
                    instructions=[
                        f"Saisissez {tr('profile.firstName')}, {tr('profile.lastName')} et {tr('profile.pseudo')}.",
                        f"Ajoutez {tr('profile.phone')} si vous souhaitez un contact plus direct.",
                        f"Utilisez l'aide du champ {tr('profile.pseudo')} pour choisir un identifiant clair.",
                    ],
                    screenshot_path=None,
                    placeholder="[SCREENSHOT: Formulaire de profil avec Prenom, Nom, Pseudo et Telephone - capture manuelle requise apres authentification]",
                    caption=None,
                ),
                StepGroup(
                    heading="Enregistrer le profil",
                    instructions=[
                        f"Relisez les informations saisies.",
                        f"Cliquez sur {tr('profile.saveProfile')}.",
                        "Attendez la redirection vers le tableau de bord pour commencer a utiliser l'application.",
                    ],
                    screenshot_path=None,
                    placeholder="[SCREENSHOT: Bouton Enregistrer le profil et confirmation de creation du profil - capture manuelle requise apres authentification]",
                    caption=None,
                ),
            ],
            notes="Note - Si l'e-mail de reference est absent, reconnectez-vous avec le bon compte avant de recommencer.",
        ),
        Chapter(
            title="Tableau de bord",
            overview="Le tableau de bord rassemble vos raccourcis utiles, vos trajets a venir et vos reservations recentes dans une vue d'ensemble rapide.",
            steps=[
                StepGroup(
                    heading="Ouvrir l'accueil",
                    instructions=[
                        f"Connectez-vous puis ouvrez {tr('nav.home')}.",
                        f"Reperez les blocs {tr('home.driverTrips')} et {tr('home.bookings')} dans la zone de synthese.",
                        "Utilisez cette page comme point de depart pour vos actions les plus frequentes.",
                    ],
                    screenshot_path=SCREENSHOTS_DIR / "screenshot-home-dashboard.png",
                    placeholder=None,
                    caption="Figure 5 - Tableau de bord",
                ),
                StepGroup(
                    heading="Consulter les prochains trajets conducteurs",
                    instructions=[
                        f"Faites defiler jusqu'a {tr('home.myUpcomingDriverTrips')}.",
                        f"Utilisez {tr('common.open')} sur une carte pour afficher ses details.",
                        f"Choisissez {tr('common.seeAll')} pour ouvrir la liste complete si besoin.",
                    ],
                    screenshot_path=SCREENSHOTS_DIR / "screenshot-home-dashboard.png",
                    placeholder="[SCREENSHOT: Bloc Mes prochains trajets conducteur avec cartes de trajet - capture manuelle requise avec permission conducteur]",
                    caption="Figure 6 - Bloc des trajets conducteur sur le tableau de bord",
                ),
                StepGroup(
                    heading="Consulter les reservations a venir",
                    instructions=[
                        f"Reperez la section {tr('home.myUpcomingBookings')}.",
                        f"Cliquez sur {tr('common.open')} pour revoir une reservation precise.",
                        f"Utilisez {tr('common.seeAll')} pour afficher tout l'historique utile.",
                    ],
                    screenshot_path=SCREENSHOTS_DIR / "screenshot-home-dashboard.png",
                    placeholder="[SCREENSHOT: Bloc Mes prochaines reservations avec cartes de reservation - capture manuelle requise avec donnees passager]",
                    caption="Figure 7 - Bloc des reservations sur le tableau de bord",
                ),
            ],
            notes="Astuce - Si une section n'apparait pas, votre compte n'a peut-etre pas encore la permission correspondante ou aucun trajet n'est enregitre pour le moment.",
        ),
        Chapter(
            title="Rechercher et reserver un trajet",
            overview="Recherchez un itineraire, comparez les resultats disponibles puis confirmez votre place depuis la fiche detaillee du trajet.",
            steps=[
                StepGroup(
                    heading="Lancer une recherche",
                    instructions=[
                        f"Ouvrez {tr('shell.findTrips')}.",
                        f"Renseignez {tr('search.departureCity')}, {tr('search.arrivalCity')} et, si besoin, {tr('search.departureDay')}.",
                        f"Cliquez sur {tr('search.searchRides')} pour afficher les propositions.",
                    ],
                    screenshot_path=SCREENSHOTS_DIR / "screenshot-find-trip-form.png",
                    placeholder=None,
                    caption="Figure 8 - Formulaire de recherche de trajet",
                ),
                StepGroup(
                    heading="Comparer les resultats",
                    instructions=[
                        f"Parcourez la page {tr('search.chooseRide')}.",
                        f"Comparez les cartes grace aux informations {tr('search.seatCount')} et {tr('search.pickupDetails')}.",
                        f"Cliquez sur une carte pour ouvrir les details du trajet choisi.",
                    ],
                    screenshot_path=SCREENSHOTS_DIR / "screenshot-find-trip-results-empty.png",
                    placeholder=None,
                    caption="Figure 9 - Page de resultats de recherche",
                ),
                StepGroup(
                    heading="Consulter les details d'un trajet",
                    instructions=[
                        f"Relisez les champs {tr('trip.departure')}, {tr('trip.arrival')}, {tr('trip.distance')} et {tr('trip.seatsLeft')}.",
                        f"Verifiez aussi {tr('trip.departureAddress')} et {tr('trip.arrivalAddress')} avant de continuer.",
                        f"Reperez le conducteur et le style du trajet pour confirmer que tout vous convient.",
                    ],
                    screenshot_path=SCREENSHOTS_DIR / "screenshot-trip-details.png",
                    placeholder=None,
                    caption="Figure 10 - Details d'un trajet",
                ),
                StepGroup(
                    heading="Reserver ou contacter le conducteur",
                    instructions=[
                        f"Cliquez sur {tr('trip.confirmBooking')} pour reserver votre place.",
                        f"Utilisez {tr('trip.contactDriver')} si vous voulez d'abord echanger dans la conversation.",
                        "Attendez le message de confirmation avant de fermer la page.",
                    ],
                    screenshot_path=SCREENSHOTS_DIR / "screenshot-contact-driver-page.png",
                    placeholder=None,
                    caption="Figure 11 - Page pour contacter le conducteur",
                ),
            ],
            notes="Note - Les suggestions de villes et certains resultats dependent des services de recherche et des trajets disponibles dans votre environnement.",
        ),
        Chapter(
            title="Mes reservations",
            overview="Retrouvez vos places reservees, ouvrez les details d'un trajet en tant que passager et gerez les actions encore disponibles.",
            steps=[
                StepGroup(
                    heading="Afficher la liste des reservations",
                    instructions=[
                        f"Ouvrez {tr('shell.bookings')}.",
                        f"Parcourez les cartes sous {tr('bookings.bookedRoutes')}.",
                        f"Cliquez sur une carte pour afficher les {tr('bookings.details').lower()}.",
                    ],
                    screenshot_path=SCREENSHOTS_DIR / "screenshot-bookings-list.png",
                    placeholder=None,
                    caption="Figure 12 - Liste des reservations",
                ),
                StepGroup(
                    heading="Verifier le detail d'une reservation",
                    instructions=[
                        f"Controlez {tr('bookings.reservationStatus')} et le badge {tr('bookings.active')} ou {tr('bookings.ended')}.",
                        f"Relisez {tr('trip.departureAddress')}, {tr('trip.arrivalAddress')} et {tr('trip.distance')}.",
                        f"Verifiez aussi le nombre indique dans {tr('bookings.passengers')}.",
                    ],
                    screenshot_path=SCREENSHOTS_DIR / "screenshot-booking-details.png",
                    placeholder=None,
                    caption="Figure 13 - Details d'une reservation",
                ),
                StepGroup(
                    heading="Contacter le conducteur ou annuler",
                    instructions=[
                        f"Utilisez {tr('trip.contactDriver')} pour ouvrir la conversation liee a cette reservation.",
                        f"Si le trajet est encore actif, cliquez sur {tr('bookings.cancelReservation')}.",
                        "Attendez le message de confirmation avant de revenir a la liste.",
                    ],
                    screenshot_path=SCREENSHOTS_DIR / "screenshot-contact-driver-page.png",
                    placeholder=None,
                    caption="Figure 14 - Conversation avec le conducteur depuis un trajet",
                ),
            ],
            notes="Astuce - Verifiez toujours l'etat de la reservation avant de partir pour savoir si le trajet est encore actif ou deja termine.",
        ),
        Chapter(
            title="Proposer et gerer mes trajets",
            overview="Si votre compte peut publier des trajets, utilisez cette zone pour creer une annonce, suivre les passagers et gerer chaque depart.",
            steps=[
                StepGroup(
                    heading="Ouvrir l'espace conducteur",
                    instructions=[
                        f"Ouvrez {tr('shell.myTrips')}.",
                        f"Parcourez {tr('driverTrips.currentTrips')}, {tr('driverTrips.incomingTrips')} et {tr('driverTrips.pastTrips')}.",
                        f"Cliquez sur {tr('driverTrips.publishNew')} pour ajouter un nouveau trajet.",
                    ],
                    screenshot_path=SCREENSHOTS_DIR / "screenshot-driver-trips-list.png",
                    placeholder=None,
                    caption="Figure 15 - Liste des trajets conducteur",
                ),
                StepGroup(
                    heading="Renseigner le trajet a publier",
                    instructions=[
                        f"Dans {tr('driverTrips.offerRide')}, completez {tr('driverTrips.dateTime')} et {tr('driverTrips.availableSeats')}.",
                        f"Activez ou laissez desactive {tr('driverTrips.smokingToggle')} selon votre preference.",
                        f"Verifiez le resume avant de passer aux adresses.",
                    ],
                    screenshot_path=SCREENSHOTS_DIR / "screenshot-driver-publish-trip.png",
                    placeholder=None,
                    caption="Figure 16 - Formulaire de publication d'un trajet",
                ),
                StepGroup(
                    heading="Selectionner les adresses suggerees",
                    instructions=[
                        f"Recherchez l'adresse de depart dans {tr('driverTrips.searchStartingAddress')}.",
                        f"Recherchez ensuite l'arrivee dans {tr('driverTrips.searchArrivalAddress')}.",
                        f"Selectionnez une suggestion pour chaque zone avant de cliquer sur {tr('driverTrips.publishTrip')}.",
                    ],
                    screenshot_path=SCREENSHOTS_DIR / "screenshot-driver-publish-address-suggestions.png",
                    placeholder=None,
                    caption="Figure 17 - Suggestions d'adresses pour publier un trajet",
                ),
                StepGroup(
                    heading="Suivre les passagers et les actions du trajet",
                    instructions=[
                        f"Ouvrez un trajet publie pour afficher {tr('driverTrips.passengerList')}.",
                        f"Utilisez {tr('driverTrips.contactPassenger')} pour joindre une personne.",
                        f"Cliquez sur {tr('driverTrips.cancelTrip')} seulement si vous devez vraiment annuler le trajet.",
                    ],
                    screenshot_path=SCREENSHOTS_DIR / "screenshot-driver-trip-details.png",
                    placeholder=None,
                    caption="Figure 18 - Details d'un trajet conducteur",
                ),
            ],
            notes="Note - La publication depend de la selection d'adresses suggerees valides. Si aucune suggestion n'apparait, verifiez votre connexion reseau et votre environnement de demonstration.",
        ),
        Chapter(
            title="Messages et coordination",
            overview="La messagerie garde vos echanges dans l'application pour organiser un depart, poser une question ou confirmer un detail pratique.",
            steps=[
                StepGroup(
                    heading="Ouvrir la boite de reception",
                    instructions=[
                        f"Ouvrez {tr('shell.chats')}.",
                        f"Reperez les badges {tr('common.new')} et l'indicateur {tr('common.connected')} ou {tr('common.updating')}.",
                        f"Cliquez sur une conversation pour l'ouvrir.",
                    ],
                    screenshot_path=SCREENSHOTS_DIR / "screenshot-chat-inbox.png",
                    placeholder=None,
                    caption="Figure 19 - Boite de reception des messages",
                ),
                StepGroup(
                    heading="Envoyer un message",
                    instructions=[
                        "Ouvrez une conversation existante ou un chat depuis une fiche trajet.",
                        f"Saisissez votre texte dans le champ {tr('chat.sendTo', {'label': tr('chat.driver').lower()})}.",
                        f"Cliquez sur {tr('common.send')} pour envoyer le message.",
                    ],
                    screenshot_path=SCREENSHOTS_DIR / "screenshot-chat-conversation.png",
                    placeholder=None,
                    caption="Figure 20 - Conversation ouverte",
                ),
                StepGroup(
                    heading="Selectionner, copier ou effacer des messages",
                    instructions=[
                        f"Touchez ou cliquez sur un message selon l'aide {tr('chat.tapToSelect')} ou {tr('chat.longPressToSelect')}.",
                        f"Utilisez {tr('chat.copyMessage')} pour copier un message unique.",
                        f"Utilisez {tr('chat.clearSelectedMessages')} ou {tr('chat.clearConversation')} pour nettoyer l'historique visible sur votre compte.",
                    ],
                    screenshot_path=SCREENSHOTS_DIR / "screenshot-chat-clear-dialog.png",
                    placeholder=None,
                    caption="Figure 21 - Dialogue d'effacement d'une conversation",
                ),
            ],
            notes="Astuce - Utilisez les messages pour confirmer le lieu exact et l'heure du depart sans quitter l'application.",
        ),
        Chapter(
            title="Mon compte et mon vehicule",
            overview="Depuis votre compte, mettez a jour vos informations personnelles puis enregistrez le vehicule utilise pour vos trajets conducteur.",
            steps=[
                StepGroup(
                    heading="Mettre a jour le profil",
                    instructions=[
                        f"Ouvrez {tr('shell.account')} puis l'onglet {tr('common.profile')}.",
                        f"Modifiez {tr('profile.pseudo')}, {tr('profile.firstName')}, {tr('profile.lastName')} et {tr('profile.phone')}.",
                        f"Cliquez sur {tr('profile.saveProfile')} pour enregistrer les changements.",
                    ],
                    screenshot_path=SCREENSHOTS_DIR / "screenshot-account-profile.png",
                    placeholder=None,
                    caption="Figure 22 - Onglet Profil du compte",
                ),
                StepGroup(
                    heading="Renseigner ou modifier le vehicule",
                    instructions=[
                        f"Passez a l'onglet {tr('common.car')}.",
                        f"Choisissez {tr('car.brand')}, recherchez {tr('car.modelSearch')} puis completez {tr('car.seats')}, {tr('car.licensePlate')} et {tr('car.color')}.",
                        f"Cliquez sur {tr('car.saveCar')} pour enregistrer le vehicule.",
                    ],
                    screenshot_path=SCREENSHOTS_DIR / "screenshot-account-car.png",
                    placeholder=None,
                    caption="Figure 23 - Onglet Voiture du compte",
                ),
                StepGroup(
                    heading="Supprimer le vehicule ou le compte",
                    instructions=[
                        f"Utilisez {tr('car.removeCar')} si vous ne souhaitez plus afficher votre vehicule.",
                        f"Dans la zone sensible du profil, cliquez sur {tr('profile.deleteAccount')} seulement si vous voulez fermer le compte.",
                        "Lisez attentivement la confirmation avant de valider une suppression definitive.",
                    ],
                    screenshot_path=None,
                    placeholder="[SCREENSHOT: Zone sensible du compte avec actions Supprimer la voiture et Supprimer le compte - capture manuelle requise avec permission compte]",
                    caption=None,
                ),
            ],
            notes="Note - Conservez vos informations de contact a jour pour faciliter la coordination avec les autres membres de la communaute.",
        ),
        Chapter(
            title="FAQ et depannage",
            overview="Cette section repond aux questions courantes et vous aide a verifier rapidement les points les plus utiles avant de demander de l'aide.",
            steps=[
                StepGroup(
                    heading="Questions frequentes",
                    instructions=[
                        "Si vous ne voyez pas une page, reconnectez-vous puis verifiez que votre profil est bien complete.",
                        "Si une liste est vide, il est possible qu'aucun trajet, aucune reservation ou aucune conversation ne soit encore disponible pour votre compte.",
                        "Si un bouton reste inactif, completez les champs obligatoires affiches a l'ecran.",
                    ],
                    screenshot_path=None,
                    placeholder="[SCREENSHOT: Aucun screenshot requis pour la FAQ - verifier les cas d'usage reels pendant la recette finale]",
                    caption=None,
                ),
                StepGroup(
                    heading="Depannage rapide",
                    instructions=[
                        "Rafraichissez la page si les donnees semblent anciennes.",
                        "Verifiez votre connexion internet si les suggestions d'adresses ou les resultats ne se chargent pas.",
                        "Reconnectez-vous si vous etes renvoye vers l'ecran de connexion pendant votre navigation.",
                    ],
                    screenshot_path=None,
                    placeholder="[SCREENSHOT: Aucun screenshot requis pour le depannage - verifier les messages d'erreur propres a votre environnement]",
                    caption=None,
                ),
            ],
            notes="Astuce - Les ecrans proteges dependent souvent de votre session, de vos permissions et des donnees disponibles dans l'environnement utilise.",
        ),
        Chapter(
            title="Contact et assistance",
            overview="Si vous avez encore besoin d'aide, preparez les informations utiles afin que votre organisation puisse vous assister plus rapidement.",
            steps=[
                StepGroup(
                    heading="Preparer votre demande",
                    instructions=[
                        "Notez la page concernee et l'action que vous tentiez d'effectuer.",
                        "Gardez une capture de l'ecran ou le texte du message affiche si un probleme apparait.",
                        "Indiquez l'heure du probleme et, si possible, le trajet ou la reservation concernes.",
                    ],
                    screenshot_path=None,
                    placeholder="[SCREENSHOT: Aucun screenshot requis pour la page de support - a completer si votre organisation dispose d'un contact visible dans l'application]",
                    caption=None,
                ),
            ],
            notes="Note - Si aucun contact n'est visible dans votre deployment Covoit, adressez-vous a l'organisation ou a l'equipe qui vous a fourni l'acces a l'application.",
        ),
    ]


def en_chapters(tr: Callable[[str, dict[str, str | int] | None], str]) -> list[Chapter]:
    return [
        Chapter(
            title="Authentication and access",
            overview="Use the sign-in screens to open your Covoit space, create an account, or recover access to your profile.",
            steps=[
                StepGroup(
                    heading="Sign in",
                    instructions=[
                        f"Open the {tr('auth.signIn')} page.",
                        f"Enter {tr('auth.emailAddress')} and {tr('auth.password')}.",
                        f"Click {tr('auth.signIn')} to enter your space.",
                    ],
                    screenshot_path=SCREENSHOTS_DIR / "screenshot-authentication-1-en.png",
                    placeholder=None,
                    caption="Figure 1 - Sign-in screen",
                ),
                StepGroup(
                    heading="Create an account",
                    instructions=[
                        f"From the sign-in screen, click {tr('auth.createAccount')}.",
                        f"Enter {tr('common.email')}, {tr('common.password')}, and {tr('common.confirm')}.",
                        f"Submit with {tr('auth.createAccount')} to move to profile completion.",
                    ],
                    screenshot_path=SCREENSHOTS_DIR / "screenshot-authentication-2-en.png",
                    placeholder=None,
                    caption="Figure 2 - Create account screen",
                ),
                StepGroup(
                    heading="Request a reset link",
                    instructions=[
                        f"Open the {tr('auth.forgotPassword')} page.",
                        f"Enter your {tr('common.email')}.",
                        f"Click {tr('auth.sendResetLinkButton')} to receive instructions.",
                    ],
                    screenshot_path=SCREENSHOTS_DIR / "screenshot-authentication-3-en.png",
                    placeholder=None,
                    caption="Figure 3 - Reset link request screen",
                ),
                StepGroup(
                    heading="Set a new password",
                    instructions=[
                        f"Open the email link and check the {tr('common.email')} field.",
                        f"Enter {tr('auth.newPassword')} and {tr('auth.confirmPassword')}.",
                        f"Click {tr('auth.resetPassword')} to finish the update.",
                    ],
                    screenshot_path=SCREENSHOTS_DIR / "screenshot-authentication-4-en.png",
                    placeholder=None,
                    caption="Figure 4 - Password reset screen",
                ),
            ],
            notes="Tip - Keep the same email across your rides so your bookings and conversations stay easy to find.",
        ),
        Chapter(
            title="First-time profile completion",
            overview="After registration, complete your basic identity details to unlock protected pages and appear clearly in the community.",
            steps=[
                StepGroup(
                    heading="Check the reference email",
                    instructions=[
                        f"Open the {tr('profile.completeProfile')} screen.",
                        f"Review the {tr('common.email')} field shown for reference.",
                        "Make sure you are completing the correct account before you continue.",
                    ],
                    screenshot_path=None,
                    placeholder="[SCREENSHOT: Complete Profile screen with the Email field visible - manual capture required after authentication]",
                    caption=None,
                ),
                StepGroup(
                    heading="Fill in your identity details",
                    instructions=[
                        f"Enter {tr('profile.firstName')}, {tr('profile.lastName')}, and {tr('profile.pseudo')}.",
                        f"Add {tr('profile.phone')} if you want a more direct contact option.",
                        f"Use the helper text around {tr('profile.pseudo')} to choose a clear username.",
                    ],
                    screenshot_path=None,
                    placeholder="[SCREENSHOT: Profile form with First Name, Last Name, Username, and Phone - manual capture required after authentication]",
                    caption=None,
                ),
                StepGroup(
                    heading="Save the profile",
                    instructions=[
                        "Review the information you entered.",
                        f"Click {tr('profile.saveProfile')}.",
                        "Wait for the redirect to the dashboard before you continue.",
                    ],
                    screenshot_path=None,
                    placeholder="[SCREENSHOT: Save Profile action and confirmation state - manual capture required after authentication]",
                    caption=None,
                ),
            ],
            notes="Note - If the reference email is missing, sign in again with the correct account before you restart the setup.",
        ),
        Chapter(
            title="Home dashboard",
            overview="The dashboard gathers your useful shortcuts, upcoming trips, and recent bookings in one quick overview.",
            steps=[
                StepGroup(
                    heading="Open the home page",
                    instructions=[
                        f"Sign in and open {tr('nav.home')}.",
                        f"Look for the {tr('home.driverTrips')} and {tr('home.bookings')} summary blocks.",
                        "Use this page as your starting point for common actions.",
                    ],
                    screenshot_path=SCREENSHOTS_DIR / "screenshot-home-dashboard.png",
                    placeholder=None,
                    caption="Figure 5 - Home dashboard",
                ),
                StepGroup(
                    heading="Review upcoming driver trips",
                    instructions=[
                        f"Scroll to {tr('home.myUpcomingDriverTrips')}.",
                        f"Use {tr('common.open')} on a card to see its details.",
                        f"Choose {tr('common.seeAll')} when you want the full list.",
                    ],
                    screenshot_path=SCREENSHOTS_DIR / "screenshot-home-dashboard.png",
                    placeholder="[SCREENSHOT: Upcoming driver trips section with trip cards - manual capture required with driver permission]",
                    caption="Figure 6 - Driver trips block on the dashboard",
                ),
                StepGroup(
                    heading="Review upcoming bookings",
                    instructions=[
                        f"Find the {tr('home.myUpcomingBookings')} section.",
                        f"Click {tr('common.open')} to revisit one booking.",
                        f"Use {tr('common.seeAll')} to open the full list.",
                    ],
                    screenshot_path=SCREENSHOTS_DIR / "screenshot-home-dashboard.png",
                    placeholder="[SCREENSHOT: Upcoming bookings section with booking cards - manual capture required with passenger data]",
                    caption="Figure 7 - Bookings block on the dashboard",
                ),
            ],
            notes="Tip - If a section is missing, your account may not have that permission yet or there may be no data to show right now.",
        ),
        Chapter(
            title="Find and book a ride",
            overview="Search for a route, compare available results, and confirm your seat from the trip details page.",
            steps=[
                StepGroup(
                    heading="Start a search",
                    instructions=[
                        f"Open {tr('shell.findTrips')}.",
                        f"Fill in {tr('search.departureCity')}, {tr('search.arrivalCity')}, and, if needed, {tr('search.departureDay')}.",
                        f"Click {tr('search.searchRides')} to load the matching rides.",
                    ],
                    screenshot_path=SCREENSHOTS_DIR / "screenshot-find-trip-form.png",
                    placeholder=None,
                    caption="Figure 8 - Ride search form",
                ),
                StepGroup(
                    heading="Compare the results",
                    instructions=[
                        f"Browse the {tr('search.chooseRide')} page.",
                        f"Compare cards using {tr('search.seatCount')} and {tr('search.pickupDetails')}.",
                        "Open the card that best matches your plan.",
                    ],
                    screenshot_path=SCREENSHOTS_DIR / "screenshot-find-trip-results-empty.png",
                    placeholder=None,
                    caption="Figure 9 - Search results page",
                ),
                StepGroup(
                    heading="Review a trip",
                    instructions=[
                        f"Read {tr('trip.departure')}, {tr('trip.arrival')}, {tr('trip.distance')}, and {tr('trip.seatsLeft')}.",
                        f"Check {tr('trip.departureAddress')} and {tr('trip.arrivalAddress')} before you continue.",
                        "Confirm that the driver and ride details match your plan.",
                    ],
                    screenshot_path=SCREENSHOTS_DIR / "screenshot-trip-details.png",
                    placeholder=None,
                    caption="Figure 10 - Trip details page",
                ),
                StepGroup(
                    heading="Book or contact the driver",
                    instructions=[
                        f"Click {tr('trip.confirmBooking')} to reserve your seat.",
                        f"Use {tr('trip.contactDriver')} if you want to ask a question first.",
                        "Wait for the confirmation message before you leave the page.",
                    ],
                    screenshot_path=SCREENSHOTS_DIR / "screenshot-contact-driver-page.png",
                    placeholder=None,
                    caption="Figure 11 - Contact Driver page",
                ),
            ],
            notes="Note - City suggestions and some search results depend on external services and the rides available in your environment.",
        ),
        Chapter(
            title="My bookings",
            overview="Review your reserved seats, open the passenger view of a trip, and manage the actions that are still available.",
            steps=[
                StepGroup(
                    heading="Open the bookings list",
                    instructions=[
                        f"Open {tr('shell.bookings')}.",
                        f"Browse the cards under {tr('bookings.bookedRoutes')}.",
                        f"Open a card to view its {tr('bookings.details').lower()}.",
                    ],
                    screenshot_path=SCREENSHOTS_DIR / "screenshot-bookings-list.png",
                    placeholder=None,
                    caption="Figure 12 - Bookings list",
                ),
                StepGroup(
                    heading="Review booking details",
                    instructions=[
                        f"Check {tr('bookings.reservationStatus')} and the {tr('bookings.active')} or {tr('bookings.ended')} badge.",
                        f"Review {tr('trip.departureAddress')}, {tr('trip.arrivalAddress')}, and {tr('trip.distance')}.",
                        f"Confirm the passenger count shown in {tr('bookings.passengers')}.",
                    ],
                    screenshot_path=SCREENSHOTS_DIR / "screenshot-booking-details.png",
                    placeholder=None,
                    caption="Figure 13 - Booking details",
                ),
                StepGroup(
                    heading="Contact the driver or cancel",
                    instructions=[
                        f"Use {tr('trip.contactDriver')} to open the booking conversation.",
                        f"If the ride is still active, click {tr('bookings.cancelReservation')}.",
                        "Wait for the confirmation before you return to the list.",
                    ],
                    screenshot_path=SCREENSHOTS_DIR / "screenshot-contact-driver-page.png",
                    placeholder=None,
                    caption="Figure 14 - Contact Driver flow",
                ),
            ],
            notes="Tip - Check the booking status before you travel so you know whether the ride is still active or already completed.",
        ),
        Chapter(
            title="Offer and manage my trips",
            overview="If your account can publish trips, use this area to create a ride, follow passengers, and manage each departure.",
            steps=[
                StepGroup(
                    heading="Open the driver area",
                    instructions=[
                        f"Open {tr('shell.myTrips')}.",
                        f"Browse {tr('driverTrips.currentTrips')}, {tr('driverTrips.incomingTrips')}, and {tr('driverTrips.pastTrips')}.",
                        f"Click {tr('driverTrips.publishNew')} to add a new ride.",
                    ],
                    screenshot_path=SCREENSHOTS_DIR / "screenshot-driver-trips-list.png",
                    placeholder=None,
                    caption="Figure 15 - My Trips page",
                ),
                StepGroup(
                    heading="Fill in the trip details",
                    instructions=[
                        f"In {tr('driverTrips.offerRide')}, complete {tr('driverTrips.dateTime')} and {tr('driverTrips.availableSeats')}.",
                        f"Turn {tr('driverTrips.smokingToggle')} on or off to match your preference.",
                        "Review the summary before you move to the address fields.",
                    ],
                    screenshot_path=SCREENSHOTS_DIR / "screenshot-driver-publish-trip.png",
                    placeholder=None,
                    caption="Figure 16 - Publish Trip form",
                ),
                StepGroup(
                    heading="Select suggested addresses",
                    instructions=[
                        f"Search the departure address in {tr('driverTrips.searchStartingAddress')}.",
                        f"Search the arrival address in {tr('driverTrips.searchArrivalAddress')}.",
                        f"Pick one suggestion for each field before you click {tr('driverTrips.publishTrip')}.",
                    ],
                    screenshot_path=SCREENSHOTS_DIR / "screenshot-driver-publish-address-suggestions.png",
                    placeholder=None,
                    caption="Figure 17 - Address suggestions in Publish Trip",
                ),
                StepGroup(
                    heading="Manage passengers and trip actions",
                    instructions=[
                        f"Open a published ride to review {tr('driverTrips.passengerList')}.",
                        f"Use {tr('driverTrips.contactPassenger')} to reach a passenger.",
                        f"Click {tr('driverTrips.cancelTrip')} only if you truly need to stop the ride.",
                    ],
                    screenshot_path=SCREENSHOTS_DIR / "screenshot-driver-trip-details.png",
                    placeholder=None,
                    caption="Figure 18 - Driver trip details",
                ),
            ],
            notes="Note - Publishing depends on valid suggested addresses. If no suggestions appear, check your network connection and demo environment.",
        ),
        Chapter(
            title="Messages and trip coordination",
            overview="Keep your ride discussions inside the app so you can confirm timing, pickup details, and trip updates in one place.",
            steps=[
                StepGroup(
                    heading="Open the inbox",
                    instructions=[
                        f"Open {tr('shell.chats')}.",
                        f"Look for {tr('common.new')} badges and the {tr('common.connected')} or {tr('common.updating')} status.",
                        "Select a conversation to open it.",
                    ],
                    screenshot_path=SCREENSHOTS_DIR / "screenshot-chat-inbox.png",
                    placeholder=None,
                    caption="Figure 19 - Chat inbox",
                ),
                StepGroup(
                    heading="Send a message",
                    instructions=[
                        "Open an existing conversation or start chat from a trip page.",
                        f"Type in the field that begins with {tr('chat.sendTo', {'label': tr('chat.driver').lower()})}.",
                        f"Click {tr('common.send')} to deliver the message.",
                    ],
                    screenshot_path=SCREENSHOTS_DIR / "screenshot-chat-conversation.png",
                    placeholder=None,
                    caption="Figure 20 - Open conversation",
                ),
                StepGroup(
                    heading="Select, copy, or clear messages",
                    instructions=[
                        f"Tap or click a message using the hint {tr('chat.tapToSelect')} or {tr('chat.longPressToSelect')}.",
                        f"Use {tr('chat.copyMessage')} to copy a single message.",
                        f"Use {tr('chat.clearSelectedMessages')} or {tr('chat.clearConversation')} to clean the visible history for your account.",
                    ],
                    screenshot_path=SCREENSHOTS_DIR / "screenshot-chat-clear-dialog.png",
                    placeholder=None,
                    caption="Figure 21 - Clear conversation dialog",
                ),
            ],
            notes="Tip - Use messages to confirm the exact pickup place and departure time without leaving the app.",
        ),
        Chapter(
            title="My account and vehicle",
            overview="From your account page, update your personal information and save the vehicle you use for driver trips.",
            steps=[
                StepGroup(
                    heading="Update the profile",
                    instructions=[
                        f"Open {tr('shell.account')} and stay on the {tr('common.profile')} tab.",
                        f"Edit {tr('profile.pseudo')}, {tr('profile.firstName')}, {tr('profile.lastName')}, and {tr('profile.phone')}.",
                        f"Click {tr('profile.saveProfile')} to save your changes.",
                    ],
                    screenshot_path=SCREENSHOTS_DIR / "screenshot-account-profile.png",
                    placeholder=None,
                    caption="Figure 22 - Profile tab",
                ),
                StepGroup(
                    heading="Add or change the vehicle",
                    instructions=[
                        f"Switch to the {tr('common.car')} tab.",
                        f"Choose {tr('car.brand')}, search {tr('car.modelSearch')}, then complete {tr('car.seats')}, {tr('car.licensePlate')}, and {tr('car.color')}.",
                        f"Click {tr('car.saveCar')} to save the vehicle.",
                    ],
                    screenshot_path=SCREENSHOTS_DIR / "screenshot-account-car.png",
                    placeholder=None,
                    caption="Figure 23 - Car tab",
                ),
                StepGroup(
                    heading="Remove the vehicle or the account",
                    instructions=[
                        f"Use {tr('car.removeCar')} if you no longer want to show a vehicle.",
                        f"In the sensitive account area, click {tr('profile.deleteAccount')} only if you intend to close the account.",
                        "Read the confirmation carefully before you approve any permanent deletion.",
                    ],
                    screenshot_path=None,
                    placeholder="[SCREENSHOT: Sensitive account area with Remove Car and Delete Account actions - manual capture required with account permission]",
                    caption=None,
                ),
            ],
            notes="Note - Keep your contact information current so other members can coordinate with you more easily.",
        ),
        Chapter(
            title="FAQ and troubleshooting",
            overview="This section answers common questions and helps you check the most useful points before you ask for help.",
            steps=[
                StepGroup(
                    heading="Frequently asked questions",
                    instructions=[
                        "If you cannot see a page, sign in again and confirm that your profile is complete.",
                        "If a list is empty, there may be no rides, bookings, or conversations available for your account yet.",
                        "If a button stays disabled, complete the required fields shown on screen.",
                    ],
                    screenshot_path=None,
                    placeholder="[SCREENSHOT: No screenshot required for the FAQ - verify real examples during final review]",
                    caption=None,
                ),
                StepGroup(
                    heading="Quick troubleshooting",
                    instructions=[
                        "Refresh the page if the information looks out of date.",
                        "Check your internet connection if address suggestions or results do not load.",
                        "Sign in again if the app returns you to the sign-in screen while you browse.",
                    ],
                    screenshot_path=None,
                    placeholder="[SCREENSHOT: No screenshot required for troubleshooting - verify environment-specific error states manually]",
                    caption=None,
                ),
            ],
            notes="Tip - Protected screens often depend on your session, your permissions, and the data available in the environment you are using.",
        ),
        Chapter(
            title="Contact and support",
            overview="If you still need help, prepare the right details so your organization can assist you faster.",
            steps=[
                StepGroup(
                    heading="Prepare your request",
                    instructions=[
                        "Note the page and the action you were trying to use.",
                        "Keep a screenshot or the visible message text if a problem appears.",
                        "Share the time of the issue and, if possible, the ride or booking involved.",
                    ],
                    screenshot_path=None,
                    placeholder="[SCREENSHOT: No screenshot required for the support page - add one only if your deployment shows a visible contact option]",
                    caption=None,
                ),
            ],
            notes="Note - If no contact is visible in your Covoit deployment, reach out to the organization or team that provided your app access.",
        ),
    ]


def build_guide(locale: str, translations: dict[str, str], brand: BrandTokens, version: str) -> dict[str, object]:
    tr = tr_factory(translations)
    document = Document()
    customize_styles(document, brand)

    build_cover(document, locale, translations["app.name"], version)
    add_toc(document, locale)

    chapters = fr_chapters(tr) if locale == "fr" else en_chapters(tr)

    screenshot_total = 0
    placeholder_total = 0
    for chapter in chapters:
        embedded, placeholders = add_chapter(document, chapter)
        screenshot_total += embedded
        placeholder_total += placeholders
        document.add_page_break()

    output_path = DOCS_DIR / ("guide-fr.docx" if locale == "fr" else "guide-en.docx")
    document.save(output_path)

    validation = validate_guide(output_path)
    return {
        "locale": locale,
        "path": str(output_path),
        "chapters": len(chapters),
        "screenshots_embedded": screenshot_total,
        "placeholders": placeholder_total,
        "validation": validation,
    }


def validate_guide(path: Path) -> dict[str, object]:
    result = {
        "opened": False,
        "toc_present": False,
        "heading1_count": 0,
        "placeholder_count": 0,
        "issues": [],
    }

    try:
        document = Document(path)
        result["opened"] = True
    except Exception as error:  # pragma: no cover - validation path only
        result["issues"].append(f"Unable to open document: {error}")
        return result

    headings = [paragraph for paragraph in document.paragraphs if paragraph.style.name == "Heading 1"]
    result["heading1_count"] = len(headings)
    if not headings:
        result["issues"].append("No Heading 1 paragraphs found.")

    with ZipFile(path) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
    result["toc_present"] = "TOC \\o \"1-3\" \\h \\z \\u" in document_xml
    if not result["toc_present"]:
        result["issues"].append("Table of contents field was not found.")

    placeholder_count = sum(1 for paragraph in document.paragraphs if "[SCREENSHOT:" in paragraph.text)
    result["placeholder_count"] = placeholder_count

    return result


def prepare_screenshots(brand: BrandTokens) -> list[str]:
    prepared: list[str] = []
    for path in sorted(SCREENSHOTS_DIR.glob("*.png")):
        ensure_bordered_image(path, brand.primary)
        prepared.append(path.name)
    return prepared


def main() -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    SCREENSHOTS_DIR.mkdir(parents=True, exist_ok=True)

    package = load_json(PACKAGE_FILE)
    dictionaries = parse_dictionaries()
    brand = parse_brand_tokens()
    prepared_screenshots = prepare_screenshots(brand)

    reports = [
        build_guide("fr", dictionaries["fr"], brand, package["version"]),
        build_guide("en", dictionaries["en"], brand, package["version"]),
    ]

    summary = {
        "brand": {
            "primary": brand.primary,
            "primary_soft": brand.primary_soft,
            "ink": brand.ink,
            "body_font": brand.body_font,
            "heading_font": brand.heading_font,
        },
        "supported_locales": ["en", "fr", "ar"],
        "prepared_screenshots": prepared_screenshots,
        "reports": reports,
        "manual_review": [
            "Replace screenshot placeholders for authenticated, permission-gated, and data-driven pages after connecting to a seeded backend.",
            "Update the table of contents in Word once the final page numbers are refreshed.",
            "Review support wording if your deployment exposes a specific contact email, phone number, or help page.",
        ],
    }

    print(json.dumps(summary, ensure_ascii=True, indent=2))


if __name__ == "__main__":
    main()
