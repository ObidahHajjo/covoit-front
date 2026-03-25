from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
DOCX_PATH = ROOT / "documentation-technique-v1.1.docx"
PDF_PATH = ROOT / "documentation-technique-v1.1.pdf"

PRIMARY = RGBColor(79, 102, 69)
ACCENT = RGBColor(198, 218, 156)
TEXT = RGBColor(45, 52, 54)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.color.rgb = RGBColor(255, 255, 255) if bold else TEXT
    run.font.size = Pt(10.5)


def add_title(document: Document, title: str, subtitle: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = PRIMARY

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(12)
    run.font.color.rgb = TEXT

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Version 1.1 - {date.today().strftime('%d/%m/%Y')}")
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT


def add_heading(document: Document, text: str, level: int = 1) -> None:
    p = document.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    run.font.color.rgb = PRIMARY


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        shade_cell(header_cells[index], "C6DA9C")
        set_cell_text(header_cells[index], header, bold=True)

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            p = cells[index].paragraphs[0]
            run = p.add_run(value)
            run.font.size = Pt(10.5)
            run.font.color.rgb = TEXT
    document.add_paragraph()


def build_document() -> Document:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].font.color.rgb = TEXT

    for style_name, size in (("Heading 1", 18), ("Heading 2", 14)):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = PRIMARY

    if "CodeBlock" not in styles:
        code_style = styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = "Consolas"
        code_style.font.size = Pt(9.5)

    add_title(document, "Covoit Frontend", "Documentation technique mise a jour")

    add_heading(document, "1. Objet du document")
    document.add_paragraph(
        "Cette version 1.1 remplace les points devenus obsoletes dans la documentation technique precedente. "
        "Elle aligne le document avec l'etat reel du frontend React/Vite et du backend Laravel JWT au 25/03/2026."
    )

    add_heading(document, "2. Stack technique actuelle")
    add_table(
        document,
        ["Element", "Valeur actuelle", "Source"],
        [
            ["React", "React 19", "package.json"],
            ["Routeur", "react-router-dom 7", "package.json"],
            ["Build", "Vite 7", "package.json"],
            ["HTTP", "Axios 1", "package.json"],
            [
                "Temps reel",
                "Laravel Echo + pusher-js + Reverb",
                "package.json / backend",
            ],
            ["i18n", "Implementation custom", "src/i18n"],
        ],
    )

    add_heading(document, "3. Authentification et session")
    document.add_paragraph(
        "Le systeme utilise des cookies httpOnly pour la session. Les reponses JSON d'authentification ne renvoient plus "
        "les jetons. Le frontend s'appuie sur les cookies `access_token` et `refresh_token` transmis automatiquement par le navigateur."
    )
    add_bullets(
        document,
        [
            "POST /auth/login : ouvre la session et definit les cookies, puis renvoie uniquement un message de succes.",
            "POST /auth/register : cree le compte, definit les cookies de session, puis renvoie uniquement un message de succes.",
            "POST /auth/refresh : renouvelle la session a partir du cookie de refresh, puis renvoie uniquement un message de succes.",
            "Le frontend restaure l'utilisateur via GET /auth/me apres login ou inscription.",
        ],
    )

    add_heading(document, "4. Contrats de reponse TypeScript")
    document.add_paragraph("Le type generique standard reste volontairement minimal.")
    p = document.add_paragraph(style="CodeBlock")
    p.add_run("export interface ApiResponse<T> {\n  data: T;\n}")
    document.add_paragraph(
        "La documentation ne doit plus presenter `message?: string` comme faisant partie du contrat global `ApiResponse<T>`. "
        "Les messages explicites sont reserves aux endpoints qui les renvoient en plus de leur enveloppe ou hors enveloppe standard."
    )

    add_heading(document, "5. Routage et garde PermissionRoute")
    document.add_paragraph(
        "La garde `PermissionRoute` a ete corrigee pour utiliser `useLocation()` de React Router au lieu de l'objet global `window.location` "
        "dans l'etat de redirection. Cette correction evite les erreurs de clonage (`DataCloneError`) lors des redirections vers /login."
    )

    add_heading(document, "6. Support WebSocket en direct")
    add_table(
        document,
        ["Canal", "Abonne", "Evenements"],
        [
            ["private-chat.user.{personId}", "ChatInboxProvider", ".chat.message.sent"],
            [
                "private-chat.conversation.{id}",
                "useChatConversation",
                ".chat.message.sent",
            ],
            [
                "private-support.session.{id}",
                "useSupportChatRealtime",
                ".support.message.sent, .support.typing",
            ],
            [
                "private-support.admins",
                "useAdminSupportRealtime",
                ".support.session.created, .support.message.sent",
            ],
        ],
    )
    document.add_paragraph(
        "Les anciens noms de canaux `chat.support.session.*` et `admin.support.inbox` ne doivent plus etre utilises dans la documentation."
    )

    add_heading(document, "7. Visibilite vehicule cote passager")
    document.add_paragraph(
        "Les passagers voient maintenant les informations vehicule du conducteur dans les ecrans de detail trajet et detail reservation."
    )
    add_bullets(
        document,
        [
            "Marque",
            "Modele",
            "Plaque d'immatriculation",
            "Couleur",
            "Pastille de couleur rendue a partir de `hex_code`",
        ],
    )
    document.add_paragraph(
        "Le backend expose ces donnees dans `TripResource` via `driver.car`, `driver.car.model.brand` et `driver.car.color`."
    )

    add_heading(document, "8. API reelle de FloatingToast")
    document.add_paragraph("Le composant `FloatingToast` accepte actuellement :")
    add_bullets(
        document,
        [
            "`message?: ReactNode | null`",
            '`tone: "success" | "error"`',
            "`durationMs?: number`",
            "Une duree plafonnee a 5000 ms",
        ],
    )
    document.add_paragraph(
        "Le ton `info` ne fait pas partie de l'implementation actuelle et ne doit pas etre documente comme supporte."
    )

    add_heading(document, "9. Security notes mises a jour")
    add_bullets(
        document,
        [
            "Aucun token d'authentification n'est persiste en localStorage ou sessionStorage cote frontend.",
            "Les cookies httpOnly restent la source de verite pour l'authentification navigateur.",
            "Les reponses JSON d'auth ne renvoient plus `access_token` ni `refresh_token`.",
            "Les guards frontend restent des protections UX ; l'autorisation reelle est appliquee par le backend.",
        ],
    )

    add_heading(document, "10. Fichiers alignes par cette mise a jour")
    add_bullets(
        document,
        [
            "Frontend : `src/features/auth/authApi.ts`, `src/hooks/Auth/useRegister.ts`, `src/router/PermissionRoute.tsx`",
            "Frontend : `src/components/ui/Trips/TripDetailsSection.tsx`, `src/components/ui/Booking/BookingDetailsSection.tsx`",
            "Backend : `app/Http/Controllers/AuthController.php`, `app/Http/Resources/AuthSessionResource.php`",
            "Backend : `app/Http/Resources/TripResource.php`, `app/Http/Controllers/TripController.php`",
        ],
    )

    document.add_paragraph(
        "Fin du document - cette version corrige les divergences majeures identifiees lors de l'audit technique."
    )
    return document


if __name__ == "__main__":
    doc = build_document()
    doc.save(DOCX_PATH)
    print(DOCX_PATH)
